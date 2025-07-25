import requests
from bs4 import BeautifulSoup
from docx import Document

url = "https://www.pcb.com.pk/about-pcb.html"
headers = {
    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36'
}

r = requests.get(url)
soup = BeautifulSoup(r.text, "html.parser")

p_tags = soup.find_all('p', class_='x_MsoNoSpacing')

# Create a Word document
doc = Document()
doc.add_heading('PCB About Page Content', 0)

# Add each paragraph to the document
for p in p_tags:
    text = p.get_text(strip=True)
    if text:
        doc.add_paragraph(text)

# Save the Word file
doc.save("pcb_about_data.docx")

print("Data saved to pcb_about_data.docx")
